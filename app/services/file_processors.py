import fitz  # PyMuPDF
from docx import Document
from PIL import Image
import pytesseract
import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)

class FileProcessor:
    """Handle file processing for different formats"""
    
    async def extract_text(self, content: bytes, file_extension: str) -> str:
        """Extract text from file content based on file type"""
        logger.info(f"Starting text extraction for file type: {file_extension}, size: {len(content)} bytes")
        
        try:
            if file_extension == 'pdf':
                logger.info("Processing PDF file...")
                return await self._extract_pdf_text(content)
            elif file_extension == 'docx':
                logger.info("Processing DOCX file...")
                return await self._extract_docx_text(content)
            elif file_extension in ['jpg', 'jpeg', 'png']:
                logger.info("Processing image file with OCR...")
                return await self._extract_image_text(content)
            elif file_extension == 'txt':
                logger.info("Processing text file...")
                try:
                    # Try UTF-8 first
                    text = content.decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        # Try latin-1 as fallback
                        text = content.decode('latin-1')
                        logger.info("Used latin-1 encoding for text file")
                    except UnicodeDecodeError:
                        try:
                            # Try cp1252 as another fallback
                            text = content.decode('cp1252')
                            logger.info("Used cp1252 encoding for text file")
                        except UnicodeDecodeError:
                            # Final fallback with error handling
                            text = content.decode('utf-8', errors='ignore')
                            logger.warning("Used UTF-8 with ignored errors for text file")
                
                logger.info(f"Extracted {len(text)} characters from text file")
                return text
            else:
                logger.error(f"Unsupported file type: {file_extension}")
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_extension} file: {str(e)}")
            raise Exception(f"Failed to extract text: {str(e)}")
    
    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF using PyMuPDF"""
        try:
            logger.info("Attempting PDF text extraction with PyMuPDF...")
            pdf_document = fitz.open(stream=content, filetype="pdf")
            text = ""
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                page_text = page.get_text()
                text += page_text + "\n"
                logger.debug(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                
            pdf_document.close()
            
            if not text.strip():
                logger.warning("No text found in PDF file using PyMuPDF")
                raise Exception("No text found in PDF file")
            
            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting PDF text with PyMuPDF: {str(e)}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            logger.info("Attempting DOCX text extraction...")
            doc = Document(io.BytesIO(content))
            text = ""
            
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    paragraph_count += 1
            
            # Also extract text from tables
            table_count = 0
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
                table_count += 1
            
            logger.info(f"Extracted text from {paragraph_count} paragraphs and {table_count} tables")
            
            if not text.strip():
                logger.warning("No text found in DOCX file")
                raise Exception("No text found in DOCX file")
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    async def _extract_image_text(self, content: bytes) -> str:
        """Extract text from image using OCR"""
        try:
            logger.info("Attempting OCR text extraction from image...")
            # Open image from bytes
            image = Image.open(io.BytesIO(content))
            logger.info(f"Image opened successfully. Size: {image.size}, Mode: {image.mode}")
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                logger.info(f"Converting image from {image.mode} to RGB")
                image = image.convert('RGB')
            
            # Use Tesseract OCR to extract text
            text = pytesseract.image_to_string(image, lang='eng')
            logger.info(f"OCR extraction completed. Text length: {len(text)} characters")
            
            if not text.strip():
                logger.warning("No text found in image file using OCR")
                raise Exception("No text found in image file")
            
            logger.info(f"Successfully extracted {len(text)} characters from image using OCR")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting image text: {str(e)}")
            raise Exception(f"Failed to extract text from image: {str(e)}")
    
    def validate_file_type(self, filename: str) -> bool:
        """Validate if file type is supported"""
        if not filename:
            return False
            
        file_extension = filename.split('.')[-1].lower()
        supported_types = ['pdf', 'docx', 'jpg', 'jpeg', 'png']
        
        return file_extension in supported_types
    
    def validate_file_size(self, file_size: int, max_size: int = 4194304) -> bool:
        """Validate file size (default 4MB limit for Vercel)"""
        return file_size <= max_size
